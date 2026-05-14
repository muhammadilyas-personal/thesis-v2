import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 1. Load Data
df = pd.read_csv("outputs/tdd_full_dataset.csv")

# 2. Aggregation for Table 1
table1_data = df.groupby("Mode").agg({
    "Time (min)": "mean",
    "Churn": "mean",
    "Tests": "mean",
    "Assert": "mean",
    "Iter": "mean"
}).reindex(["Manual", "AI-TDD", "Hybrid"]).round(1).reset_index()

table1_data.columns = ["Mode", "Mean Time (min)", "LOC/Churn", "Mean Tests", "Mean Assertions", "Mean Iterations"]

# 3. Aggregation for Table 2
table2_data = df.groupby("Mode").agg({
    "Coverage %": "mean",
    "Mutation Score %": "mean",
    "Defects": "mean",
    "Cyclomatic": "mean",
    "Cog Load": "mean",
    "Churn": "mean"
}).reindex(["Manual", "AI-TDD", "Hybrid"]).round(1).reset_index()

table2_data.columns = ["Mode", "Coverage %", "Mutation Score %", "Defect Density", "Cyclomatic Complexity", "Cognitive Load", "Code Churn"]

# --- GENERATE WORD DOCUMENT ---
doc = Document()

def add_table_with_caption(doc, data, caption, observation):
    # Add Caption
    p_cap = doc.add_paragraph()
    run_cap = p_cap.add_run(caption)
    run_cap.italic = True
    p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add Table
    table = doc.add_table(rows=1, cols=len(data.columns))
    table.style = 'Table Grid'
    
    # Header Row
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(data.columns):
        hdr_cells[i].text = col_name
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data Rows
    for _, row in data.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            
    # Add Observation
    p_obs = doc.add_paragraph()
    p_obs.add_run("\nObservation: ").bold = True
    p_obs.add_run(observation)
    doc.add_paragraph() # Spacer

# Table 1 Observations
obs1 = (f"AI-TDD reduces development time by {((table1_data.iloc[0,1]-table1_data.iloc[1,1])/table1_data.iloc[0,1]*100):.1f}% "
        f"compared to Manual TDD but results in the fewest mean iterations ({table1_data.iloc[1,5]}). "
        f"Hybrid mode balances high assertion counts ({table1_data.iloc[2,4]}) with a moderate time efficiency of "
        f"{((table1_data.iloc[0,1]-table1_data.iloc[2,1])/table1_data.iloc[0,1]*100):.1f}% reduction.")

# Table 2 Observations
obs2 = (f"While AI-TDD achieves the highest coverage ({table2_data.iloc[1,1]}%), its mutation score ({table2_data.iloc[1,2]}%) "
        f"is lower than both Manual and Hybrid modes, indicating shallower test quality. "
        f"Hybrid AI-Assisted TDD maximizes quality metrics, achieving a peak mutation score of {table2_data.iloc[2,2]}% "
        f"while maintaining significantly lower cognitive load than Manual TDD.")

add_table_with_caption(doc, table1_data, "Table 1 — Classified Data (Time, LOC, Assertions)", obs1)
add_table_with_caption(doc, table2_data, "Table II — Classified Metrics Data", obs2)

doc.save("outputs/thesis_tables.docx")

# --- GENERATE MARKDOWN ---
md_content = "# Thesis Results Tables\n\n"
md_content += "### Table 1 — Classified Data (Time, LOC, Assertions)\n"
md_content += table1_data.to_markdown(index=False) + "\n\n"
md_content += "**Observation:** " + obs1 + "\n\n"

md_content += "### Table II — Classified Metrics Data\n"
md_content += table2_data.to_markdown(index=False) + "\n\n"
md_content += "**Observation:** " + obs2 + "\n"

with open("outputs/thesis_tables.md", "w") as f:
    f.write(md_content)

print("Tables generated successfully in outputs/ folder.")
