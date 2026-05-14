from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Heading
h1 = doc.add_heading('6. Conclusion and Future Work', level=1)

# 6.1 Summary of Findings
doc.add_heading('6.1 Summary of Findings', level=2)
doc.add_paragraph(
    "This research provided a rigorous empirical evaluation of the impact of AI-based Test-Driven Development (TDD) "
    "on software quality and development efficiency. The results from the study of 10 professional engineers at "
    "MSE Technology (Pakistan) across 30 experimental sessions revealed that AI integration significantly alters "
    "the TDD lifecycle. Fully AI-TDD achieved a substantial 48.50% reduction in development time compared to the "
    "Manual TDD baseline, reducing the mean task time from 26.60 to 13.70 minutes. However, this efficiency gain "
    "was accompanied by a notable quality deficit; Fully AI-TDD produced the highest defect density (\u03bc = 7.20) "
    "and a mutation score of only 38.40%, significantly lower than the Manual baseline of 48.40%."
)
doc.add_paragraph(
    "In contrast, the Hybrid AI-Assisted mode demonstrated the most balanced performance profile. It achieved a "
    "29.32% reduction in development time (\u03bc = 18.80 min) while simultaneously reaching the peak mutation "
    "score of 58.30% and the lowest defect rate of 2.20 per task. These findings suggest that the integration of "
    "human oversight into the AI workflow not only maintains but enhances the quality of test oracles while "
    "preserving significant efficiency benefits."
)

# 6.2 Answers to Research Questions
doc.add_heading('6.2 Answers to Research Questions', level=2)
doc.add_paragraph("RQ1: AI-based TDD provides a statistically significant improvement in development efficiency, with Fully AI-TDD and Hybrid modes reducing completion time by 48.5% and 29.3% respectively.")
doc.add_paragraph("RQ2: Full AI automation negatively impacts code quality by increasing defect density and producing 'shallow' test suites that fail to detect semantic mutations.")
doc.add_paragraph("RQ3: The Hybrid AI-Assisted model provides a superior trade-off, achieving higher quality metrics than Manual TDD while remaining significantly faster than traditional methods.")
doc.add_paragraph("RQ4: AI integration significantly reduces the cognitive load of developers, with the Hybrid mode providing a balanced mental workload that supports thoroughness without the fatigue of manual coding.")
doc.add_paragraph("RQ5: Human supervision is a non-negotiable requirement for maintaining the robustness of test oracles, as evidenced by the 20% mutation score gap between supervised and unsupervised AI modes.")

# 6.3 Recommendations
doc.add_heading('6.3 Recommendations', level=2)
doc.add_paragraph(
    "Based on these empirical findings, the study recommends a contextual application of AI-TDD modes. "
    "Fully AI-TDD is highly recommended for rapid prototyping, internal utility development, and scenarios where "
    "speed is the primary constraint and code longevity is minimal. For mission-critical production systems, "
    "security-sensitive modules, and complex business logic, the Hybrid AI-Assisted mode is the recommended standard, "
    "as it ensures high mutation coverage and low defect density. Furthermore, the Hybrid mode should be utilized as "
    "a mentoring tool for junior developers, allowing them to review high-speed AI drafts to improve their "
    "understanding of test structure and edge-case identification."
)

# 6.4 Limitations
doc.add_heading('6.4 Limitations', level=2)
doc.add_paragraph(
    "The findings of this study are subject to several limitations. The sample size was constrained to 10 participants "
    "(30 sessions total), which, while sufficient for observing large effect sizes, may not capture the full diversity "
    "of the global software engineering population. The study was localized to a single company, MSE Technology in "
    "Pakistan, and utilized a single task family (Login Service). Furthermore, the experiment relied on a specific "
    "family of Large Language Models (LLMs), and results may vary with different model architectures or prompt engineering "
    "frameworks."
)

# 6.5 Future Work
doc.add_heading('6.5 Future Work', level=2)
doc.add_paragraph(
    "Future research will aim to expand the participant cohort to include multiple international sites and a broader "
    "range of architectural complexities beyond micro-services. A longitudinal study will be conducted to evaluate "
    "the long-term impact of AI-TDD on code maintainability and technical debt over multiple release cycles."
)
doc.add_paragraph(
    "Additionally, future work should explore the role of Multi-Agent AI systems in TDD, specifically investigating "
    "if specialized 'Tester Agents' can close the mutation score gap observed in this study. Further investigation "
    "is also required into the educational impact of AI-TDD, specifically how the use of AI in early career "
    "development influences the fundamental TDD skills and problem-solving patterns of new software engineers."
)

doc.save("outputs/conclusion_and_future_work.docx")
print("Conclusion and Future Work section generated in outputs/ folder.")
