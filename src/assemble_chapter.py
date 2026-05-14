import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Load aggregated data for table recreation
df = pd.read_csv("outputs/tdd_full_dataset.csv")

# Agregation for Table 1
t1 = df.groupby("Mode").agg({
    "Time (min)": "mean",
    "Churn": "mean",
    "Tests": "mean",
    "Assert": "mean",
    "Iter": "mean"
}).reindex(["Manual", "AI-TDD", "Hybrid"]).round(1).reset_index()
t1.columns = ["Mode", "Mean Time (min)", "LOC/Churn", "Mean Tests", "Mean Assertions", "Mean Iterations"]

# Agregation for Table 2
t2 = df.groupby("Mode").agg({
    "Coverage %": "mean",
    "Mutation Score %": "mean",
    "Defects": "mean",
    "Cyclomatic": "mean",
    "Cog Load": "mean",
    "Churn": "mean"
}).reindex(["Manual", "AI-TDD", "Hybrid"]).round(1).reset_index()
t2.columns = ["Mode", "Coverage %", "Mutation Score %", "Defect Density", "Cyclomatic Complexity", "Cognitive Load", "Code Churn"]

# Load Statistical Text
with open("outputs/statistical_analysis_section.md", "r") as f:
    stats_content = f.read()

doc = Document()

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h

def add_caption(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table(doc, data):
    table = doc.add_table(rows=1, cols=len(data.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(data.columns):
        hdr_cells[i].text = col
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    for _, row in data.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    return table

# 5. Results and Discussion
add_heading(doc, '5. Results and Discussion', 1)

# 5.1 Dataset Overview
add_heading(doc, '5.1 Dataset Overview', 2)
doc.add_paragraph(
    "The empirical foundation of this study is a controlled experiment involving 10 professional software engineers "
    "from MSE Technology (Pakistan). The participants, representing a cross-section of experience levels (Beginner to Expert) "
    "and roles (QA, Full Stack, and Software Engineers), completed a standard 'Login Service' TDD task across three "
    "distinct experimental modes. This resulted in a paired dataset of 30 experimental sessions (P1\u2013P30). "
    "For each session, 21 primary metrics were captured, including temporal efficiency, code churn, and cognitive load, "
    "supplemented by a 22nd metric\u2014Mutation Score%\u2014to evaluate the semantic depth of the generated test suites."
)

# 5.2 Classified Data \u2014 Time, LOC, Assertions
add_heading(doc, '5.2 Classified Data \u2014 Time, LOC, Assertions', 2)
doc.add_paragraph(
    "The temporal performance and development throughput across the three modes are summarized in Table 1. "
    "Initial analysis reveals a stark contrast in development speed when AI automation is introduced."
)
add_table(doc, t1)
add_caption(doc, "Table 1 \u2014 Classified Data (Time, LOC, Assertions)")

doc.add_paragraph()
doc.add_picture("outputs/figure_2_classified_data.png", width=Inches(6))
add_caption(doc, "Figure II \u2014 Grouped comparison of mean development metrics across experimental modes.")

doc.add_paragraph(
    "As shown in Figure II, the transition from Manual TDD to Fully AI-TDD results in a dramatic reduction in "
    "mean time (26.6 min vs. 13.7 min) and mean iterations (25.2 vs. 5.7). This suggests that AI serves as a high-speed "
    "accelerator for the TDD cycle, though this speed comes at the cost of significantly fewer assertions being generated. "
    "The Hybrid mode, however, demonstrates an 'optimized' profile, maintaining the high assertion density of manual "
    "coding (8.6) while achieving nearly a 30% time reduction relative to the baseline."
)

# 5.3 Classified Metrics \u2014 Quality Indicators
add_heading(doc, '5.3 Classified Metrics \u2014 Quality Indicators', 2)
doc.add_paragraph(
    "To understand the quality trade-offs inherent in these efficiency gains, Table II summarizes the core software "
    "quality and complexity metrics. These indicators provide a more granular view of the technical debt and defect "
    "density associated with varying levels of AI automation."
)
add_table(doc, t2)
add_caption(doc, "Table II \u2014 Classified Metrics Data")

doc.add_paragraph(
    "The quality data highlights a divergence between structural coverage and semantic robustness. Fully AI-TDD "
    "achieved peak code coverage (93.3%) but also the highest defect rate (7.2 per task), confirming that broad coverage "
    "does not equate to correct logic. Conversely, the Hybrid mode achieved the lowest defect rate (2.2) and the "
    "lowest cognitive load (3.9), suggesting that developer-led AI assistance optimizes for both quality and mental "
    "well-being."
)

# 5.4 Comparative Quality\u2013Efficiency Analysis
add_heading(doc, '5.4 Comparative Quality\u2013Efficiency Analysis', 2)
doc.add_paragraph()
doc.add_picture("outputs/figure_3_quality_efficiency_matrix.png", width=Inches(5))
add_caption(doc, "Figure III \u2014 Spider plot illustrating the trade-off matrix between speed, quality, and cognitive load.")

doc.add_paragraph(
    "Figure III provides a holistic visualization of the trade-off triangle. The 'Manual' polygon is characterized by "
    "high cognitive load and low speed, whereas the 'AI-TDD' polygon expands aggressively toward speed and coverage but "
    "contracts on mutation score and defect-free code. The 'Hybrid' polygon represents the most balanced multi-dimensional "
    "performance, occupying the largest overall area within the matrix."
)
doc.add_paragraph(
    "This visualization reinforces the argument that peak software quality is not a byproduct of full automation, "
    "but rather a result of an augmented workflow where AI handles breadth (coverage) and humans handle depth (mutation/oracles). "
    "This finding is critical for organizations considering the full displacement of manual testing with AI-driven pipelines."
)

# 5.5 Statistical Analysis
add_heading(doc, '5.5 Statistical Analysis', 2)
# Re-insert the stats analysis content formatted
for block in stats_content.split('\n\n'):
    if block.strip():
        p = doc.add_paragraph(block.strip())
        if block.strip().startswith(('I.', 'II.', 'III.', 'IV.', 'V.', 'VI.', 'VII.', 'STATISTICAL')):
            p.runs[0].bold = True

# 5.6 Mapping Results to Research Questions
add_heading(doc, '5.6 Mapping Results to Research Questions', 2)

doc.add_paragraph(
    "RQ1: Does AI-based TDD significantly improve development efficiency? Yes, Fully AI-TDD reduced development time "
    "by 48.5% and Hybrid mode by 29.32% over Manual TDD. These findings provide strong empirical support for RQ1, "
    "aligning with the productivity gains documented by Dakhel et al. (2023) [2] and Wang et al. (2024) [7]."
).bold = False

doc.add_paragraph(
    "RQ2: What is the impact of full AI automation on code quality and defect rates? Full automation results in a "
    "quality decline, with defects doubling from 3.6 to 7.2 per task. As observed by Yang et al. (2024) [8], AI often "
    "prioritizes superficial coverage over semantic correctness, leading to the higher defect rates seen in our dataset."
)

doc.add_paragraph(
    "RQ3: Can a Hybrid AI-Assisted model provide a superior trade-off between speed and quality? Yes, Hybrid TDD "
    "achieved the highest mutation score (58.3%) and lowest defect density (2.2) while remaining 30% faster than manual. "
    "This supports the hybrid collaborative model advocated in Santos et al. (2023) [6]."
)

doc.add_paragraph(
    "RQ4: How does AI integration influence the cognitive load of developers during TDD? Hybrid TDD reduced mean "
    "cognitive load from 5.0 (Manual) to 3.9, while Fully AI-TDD dropped to 2.0. This suggests that AI acts as a "
    "cognitive offloading mechanism, matching the 'AI as a Co-pilot' paradigm described by Mock et al. (2024) [4]."
)

doc.add_paragraph(
    "RQ5: Is human supervision required to maintain test oracle robustness in AI-TDD? Yes, the mutation score "
    "gap between Fully AI-TDD (38.4%) and Hybrid TDD (58.3%) demonstrates that human oversight is essential for oracle "
    "robustness. This confirms the 'Human-in-the-Loop' necessity highlighted by Zhang et al. (2025) [9] and the SWEBOK [10]."
)

# 5.7 Threats to Validity
add_heading(doc, '5.7 Threats to Validity', 2)
doc.add_paragraph(
    "Internal Validity: The paired experimental design controlled for individual participant variance. However, a potential "
    "threat is the 'carry-over effect,' where learning from the manual session might have improved performance in subsequent "
    "AI modes. To mitigate this, future studies should employ a randomized Latin Square design."
)
doc.add_paragraph(
    "External Validity: The study was conducted at a single site (MSE Technology, Pakistan) using a single task (Login Service). "
    "While the participants represent a diverse skill set, the results may not generalize to all software domains or "
    "more complex, multi-component architectural tasks as defined in Santos et al. (2023) [6]."
)
doc.add_paragraph(
    "Construct Validity: The metrics used (e.g., Churn as a proxy for LOC, Mutation Score for quality) are standard in literature "
    "but may not capture the full nuance of software maintainability. The use of 'Defects' as a raw count rather than "
    "severity-weighted density is an additional limitation."
)
doc.add_paragraph(
    "Conclusion Validity: With a sample size of n=30 across three modes, the statistical power is sufficient for the "
    "observed large effect sizes (Cohen's d > 2.0). However, the variance in AI performance indicates that a larger "
    "longitudinal study would be required to confirm these trends across entire development cycles."
)

# 5.8 Discussion Summary
add_heading(doc, '5.8 Discussion Summary', 2)
doc.add_paragraph(
    "In summary, Chapter 5 demonstrates that while AI-TDD provides an unprecedented acceleration of the TDD cycle, "
    "it creates a quality deficit characterized by shallow tests and logic errors. The findings advocate for a "
    "Hybrid AI-Assisted model, which optimizes the speed-quality frontier. This empirical evidence validates the thesis "
    "hypothesis that the future of TDD lies not in total automation, but in a supervised collaborative workflow that "
    "leverages AI for breadth and humans for depth."
)

doc.save("outputs/results_and_discussion_chapter.docx")
print("Full chapter assembled successfully in outputs/ folder.")
